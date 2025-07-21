import os
import json
import logging
import gradio as gr
from pathlib import Path
from typing import List, Tuple, Dict
import time

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
# from langchain_community.embeddings import DashScopeEmbeddings  # 不再使用
from langchain_huggingface import HuggingFaceEmbeddings
from transformers import AutoTokenizer, AutoModelForCausalLM
from transformers import pipeline
from langchain_huggingface import HuggingFacePipeline
from langchain.schema import Document
from docx import Document as DocxDocument

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class EnhancedRAGSystem:
    def __init__(self):
        """初始化增强版RAG系统"""
        self.embeddings = None
        self.llm = None
        self.vector_store = None
        self.chain = None
        self.documents = []
        
        # 配置路径
        self.files_dir = Path("/root/autodl-tmp/RAG检索/file_t")
        self.files_dir.mkdir(exist_ok=True)
        
        # 微调模型路径
        self.model_path = "/root/autodl-tmp/LLaMA-Factory/output/deepseekr1_7b_lora_sft_dpo"
        
        # 嵌入模型路径
        self.embedding_model_path = "/root/autodl-tmp/Hugging-Face/models/bge-large-zh"
        
        logger.info("增强版RAG系统初始化完成")
    
    def load_models(self):
        """加载所需的模型"""
        try:
            logger.info("开始加载模型...")
            
            # 加载嵌入模型 - 使用DashScope
            # self.embeddings = DashScopeEmbeddings(
            #     model="text-embedding-v2",
            #     dashscope_api_key="sk-8f7884619fd6405f9c90fe16733c597e"
            # )
            
            # 加载嵌入模型 - 使用本地bge-large-zh
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.embedding_model_path,
                model_kwargs={'device': 'cuda'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info("嵌入模型加载完成")
            
            # 加载微调的自定义模型
            logger.info("开始加载微调模型...")
            tokenizer = AutoTokenizer.from_pretrained(self.model_path, trust_remote_code=True)
            model = AutoModelForCausalLM.from_pretrained(
                self.model_path,
                device_map="auto",
                trust_remote_code=True
            ).eval()
            
            llm_pipeline = pipeline(
                "text-generation",
                model=model,
                tokenizer=tokenizer,
                max_new_tokens=512,
                temperature=0.7,
                top_p=0.9,
                top_k=50,
                repetition_penalty=1.1,
                do_sample=True,
                return_full_text=False,
                pad_token_id=tokenizer.eos_token_id
            )
            self.llm = HuggingFacePipeline(pipeline=llm_pipeline)
            logger.info("微调模型加载完成")
            
            return "✅ 所有模型加载成功"
            
        except Exception as e:
            error_msg = f"❌ 模型加载失败: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def read_txt_file(self, file_path: Path) -> str:
        """读取txt文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
    
    def read_docx_file(self, file_path: Path) -> str:
        """读取docx文件"""
        doc = DocxDocument(file_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        return '\n'.join(content)
    
    def read_json_file(self, file_path: Path) -> str:
        """读取JSON文件并提取文本内容，支持JSONL格式"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                
            # 检查是否为JSONL格式（每行一个JSON对象）
            if '\n' in content and not content.startswith('['):
                # JSONL格式处理
                texts = []
                for line in content.split('\n'):
                    line = line.strip()
                    if line:
                        try:
                            data = json.loads(line)
                            texts.append(self._extract_text_from_json(data))
                        except json.JSONDecodeError as e:
                            logger.warning(f"跳过无效JSON行: {line[:50]}... 错误: {e}")
                            continue
                return '\n'.join(filter(None, texts))
            else:
                # 标准JSON格式处理
                data = json.loads(content)
                return self._extract_text_from_json(data)
                
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read().strip()
                
            # 检查是否为JSONL格式
            if '\n' in content and not content.startswith('['):
                texts = []
                for line in content.split('\n'):
                    line = line.strip()
                    if line:
                        try:
                            data = json.loads(line)
                            texts.append(self._extract_text_from_json(data))
                        except json.JSONDecodeError as e:
                            logger.warning(f"跳过无效JSON行: {line[:50]}... 错误: {e}")
                            continue
                return '\n'.join(filter(None, texts))
            else:
                data = json.loads(content)
                return self._extract_text_from_json(data)
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析错误: {e}")
            return ""
    
    def _extract_text_from_json(self, data) -> str:
        """从JSON数据中递归提取文本内容"""
        texts = []
        
        if isinstance(data, list):
            for item in data:
                texts.append(self._extract_text_from_json(item))
        elif isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, str) and value.strip():
                    texts.append(value.strip())
                elif isinstance(value, (dict, list)):
                    texts.append(self._extract_text_from_json(value))
                elif value is not None:
                    texts.append(str(value))
        elif isinstance(data, str) and data.strip():
            texts.append(data.strip())
        elif data is not None:
            texts.append(str(data))
        
        return '\n'.join(filter(None, texts))
    
    def load_documents(self):
        """加载所有文档"""
        try:
            self.documents = []
            supported_extensions = {'.txt', '.docx', '.json'}
            
            files_found = 0
            for file_path in self.files_dir.iterdir():
                if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                    try:
                        if file_path.suffix.lower() == '.txt':
                            content = self.read_txt_file(file_path)
                        elif file_path.suffix.lower() == '.docx':
                            content = self.read_docx_file(file_path)
                        elif file_path.suffix.lower() == '.json':
                            content = self.read_json_file(file_path)
                        
                        if content.strip():
                            # 创建LangChain Document对象
                            doc = Document(
                                page_content=content.strip(),
                                metadata={'filename': file_path.name, 'source': str(file_path)}
                            )
                            self.documents.append(doc)
                            files_found += 1
                            logger.info(f"已加载文档: {file_path.name}")
                    
                    except Exception as e:
                        logger.warning(f"无法读取文件 {file_path.name}: {str(e)}")
            
            if files_found == 0:
                return "❌ 未找到任何支持的文档文件"
            
            logger.info(f"✅ 成功加载 {files_found} 个文档")
            return f"✅ 成功加载 {files_found} 个文档"
            
        except Exception as e:
            error_msg = f"❌ 加载文档失败: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def create_chunks_and_index(self):
        """创建文档块并构建向量索引"""
        try:
            if not self.documents:
                return "❌ 没有文档，请先加载文档"
            
            if self.embeddings is None:
                return "❌ 嵌入模型未加载，请先加载模型"
            
            logger.info(f"开始处理 {len(self.documents)} 个文档")
            
            # 使用LangChain的文本分割器
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50
            )
            chunks = text_splitter.split_documents(self.documents)
            logger.info(f"文档分割完成，共 {len(chunks)} 个文档块")
            
            # 检查文档块内容
            for i, chunk in enumerate(chunks[:3]):  # 只检查前3个
                logger.info(f"文档块 {i+1}: {chunk.page_content[:100]}...")
            
            # 测试嵌入模型
            try:
                test_text = "测试文本"
                test_embedding = self.embeddings.embed_query(test_text)
                logger.info(f"嵌入模型测试成功，向量维度: {len(test_embedding)}")
            except Exception as embed_error:
                logger.error(f"嵌入模型测试失败: {embed_error}")
                return f"❌ 嵌入模型测试失败: {embed_error}"
            
            # 构建FAISS向量存储
            logger.info("开始构建FAISS向量存储...")
            self.vector_store = FAISS.from_documents(
                documents=chunks,
                embedding=self.embeddings
            )
            
            logger.info(f"✅ 创建了 {len(chunks)} 个文档块并构建索引")
            return f"✅ 创建了 {len(chunks)} 个文档块并构建索引"
            
        except Exception as e:
            error_msg = f"❌ 创建文档块和索引失败: {str(e)}"
            logger.error(error_msg)
            import traceback
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            return error_msg
    
    def build_chain(self):
        """构建问答链"""
        try:
            if self.vector_store is None:
                return "❌ 向量存储未构建，请先创建索引"
            
            if self.llm is None:
                return "❌ LLM模型未加载，请先加载模型"
            
            # 创建检索器 - 增加检索数量
            retriever = self.vector_store.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}  # 检索前5个最相关的文档
            )
            
            # 创建提示模板 - 简化格式
            template = """根据以下信息回答问题：
{context}

问题：{question}

回答："""
            
            prompt_template = ChatPromptTemplate.from_template(template)
            
            # 将lambda函数定义为命名变量
            # prompt_to_string = lambda x: x.to_string()
            
            # 定义上下文格式化函数
            def format_docs(docs):
                return "\n\n".join([doc.page_content for doc in docs])
            
            # 构建链
            self.chain = (
                RunnableParallel({"context": retriever | format_docs, "question": RunnablePassthrough()})
                | prompt_template
                # | prompt_to_string
                | self.llm
                | StrOutputParser()
            )
            
            logger.info("✅ 问答链构建完成")
            return "✅ 问答链构建完成"
            
        except Exception as e:
            error_msg = f"❌ 构建问答链失败: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def ask_question(self, question: str) -> str:
        """回答问题"""
        try:
            if not question or question.strip() == "":
                return "请输入您的问题，我将基于知识库为您提供答案。"
            
            if self.chain is None:
                return "❌ 问答链未构建，请先初始化系统"
            
            logger.info(f"开始处理问题: {question[:50]}...")
            
            # 先测试检索功能
            retriever = self.vector_store.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 5}
            )
            docs = retriever.invoke(question)
            logger.info(f"检索到 {len(docs)} 个相关文档")
            
            if docs:
                for i, doc in enumerate(docs):
                    logger.info(f"文档 {i+1}: {doc.page_content[:100]}...")
            else:
                logger.warning("未检索到相关文档")
            
            # 调用链生成答案
            answer = self.chain.invoke(question)
            logger.info(f"生成的答案: {answer[:100]}...")
            
            # 确保返回字符串
            if isinstance(answer, str):
                return answer
            else:
                return str(answer)
            
        except Exception as e:
            error_msg = f"❌ 回答问题时出错: {str(e)}"
            logger.error(error_msg)
            return error_msg


# 全局RAG系统实例
rag_system = EnhancedRAGSystem()


def initialize_system():
    """初始化整个系统"""
    steps = []
    
    # 1. 加载模型
    result = rag_system.load_models()
    steps.append(f"1. 加载模型: {result}")
    if "❌" in result:
        return "\n".join(steps)
    
    # 2. 加载文档
    result = rag_system.load_documents()
    steps.append(f"2. 加载文档: {result}")
    if "❌" in result:
        return "\n".join(steps)
    
    # 3. 创建文档块和索引
    result = rag_system.create_chunks_and_index()
    steps.append(f"3. 创建文档块和索引: {result}")
    if "❌" in result:
        return "\n".join(steps)
    
    # 4. 构建问答链
    result = rag_system.build_chain()
    steps.append(f"4. 构建问答链: {result}")
    
    return "\n".join(steps)


def process_query(query, history):
    """处理用户查询"""
    if not query.strip():
        return history, ""
    
    # 添加用户问题到历史
    history.append({"role": "user", "content": query})
    
    # 生成回答
    answer = rag_system.ask_question(query)
    
    # 添加系统回答到历史
    history.append({"role": "assistant", "content": answer})
    
    return history, ""


def format_chat_history(history):
    """格式化聊天历史"""
    if not history:
        return "暂无对话历史"
    
    formatted = []
    for msg in history:
        role = "🙋 用户" if msg["role"] == "user" else "🤖 助手"
        formatted.append(f"{role}: {msg['content']}")
    
    return "\n\n".join(formatted)


# 创建Gradio界面
def create_interface():
    with gr.Blocks(title="增强版RAG知识库问答系统", theme=gr.themes.Soft()) as interface:
        gr.Markdown("# 🚀 增强版RAG知识库问答系统")
        gr.Markdown("""
        这是一个基于LangChain和微调模型的增强版RAG知识库问答系统。
        系统支持多种文件格式（txt、docx、json），使用DashScope嵌入模型和您的微调DeepSeek-R1模型。
        """)
        
        with gr.Row():
            with gr.Column(scale=2):
                # 系统初始化部分
                gr.Markdown("## 📋 系统初始化")
                init_button = gr.Button("🔧 初始化系统", variant="primary", size="lg")
                init_status = gr.Textbox(
                    label="初始化状态",
                    lines=8,
                    interactive=False,
                    placeholder="点击上方按钮开始初始化系统..."
                )
                
                # 聊天历史显示
                gr.Markdown("## 💬 对话历史")
                chat_history_display = gr.Textbox(
                    label="聊天记录",
                    lines=10,
                    interactive=False,
                    placeholder="对话历史将在这里显示..."
                )
            
            with gr.Column(scale=3):
                # 问答交互部分
                gr.Markdown("## 🤔 智能问答")
                
                # 聊天历史状态
                chat_history = gr.State([])
                
                # 输入框和按钮
                with gr.Row():
                    query_input = gr.Textbox(
                        label="请输入您的问题",
                        placeholder="例如：宫颈口粘连怎么回事？",
                        scale=4
                    )
                    ask_button = gr.Button("🚀 提问", variant="primary", scale=1)
                
                # 清空对话按钮
                clear_button = gr.Button("🗑️ 清空对话", variant="secondary")
                
                # 示例问题
                gr.Markdown("### 💡 示例问题")
                example_questions = [
                    "宫颈口粘连怎么回事？",
                    "前列腺炎应该吃什么食物比较好？",
                    "中风有什么症状？",
                    "韧带撕裂自己会愈合吗？",
                    "卵巢囊肿微创手术多长时间恢复？"
                ]
                
                for question in example_questions:
                    example_btn = gr.Button(f"📝 {question}", variant="secondary", size="sm")
                    example_btn.click(
                        fn=lambda q=question: q,
                        outputs=query_input
                    )
        
        # 事件绑定
        init_button.click(
            fn=initialize_system,
            outputs=init_status
        )
        
        # 提问按钮点击事件
        ask_button.click(
            fn=process_query,
            inputs=[query_input, chat_history],
            outputs=[chat_history, query_input]
        ).then(
            fn=format_chat_history,
            inputs=chat_history,
            outputs=chat_history_display
        )
        
        # 回车键提问
        query_input.submit(
            fn=process_query,
            inputs=[query_input, chat_history],
            outputs=[chat_history, query_input]
        ).then(
            fn=format_chat_history,
            inputs=chat_history,
            outputs=chat_history_display
        )
        
        # 清空对话
        clear_button.click(
            fn=lambda: ([], ""),
            outputs=[chat_history, chat_history_display]
        )
        
        # 添加使用说明
        gr.Markdown("""
        ## 📖 使用说明
        
        1. **初始化系统**: 首次使用需要点击"初始化系统"按钮，系统会自动加载模型和文档
        2. **提问**: 在输入框中输入问题，点击"提问"按钮或按回车键
        3. **查看历史**: 右侧会显示完整的对话历史
        4. **示例问题**: 可以点击示例问题快速体验
        
        ## ⚙️ 系统特点
        
        - 基于LangChain框架，架构更加标准化
        - 使用DashScope嵌入模型和bge-large-zh嵌入模型
        - 集成您的微调DeepSeek-R1模型
        - 支持多种文件格式：txt、docx、json
        - 完全本地化部署，数据安全
        
        ## 🔧 自定义文档
        
        您可以将自己的文档文件放入`files`文件夹，然后重新初始化系统。
        
        **支持的文件格式：**
        - `.txt` 文件：纯文本文档
        - `.docx` 文件：Word文档  
        - `.json` 文件：JSON格式数据（如test_encyclopedia.json）
        """)
    
    return interface


if __name__ == "__main__":
    # 启动界面
    interface = create_interface()
    interface.launch(
        server_name="0.0.0.0",
        server_port=7862,  # 使用不同的端口避免冲突
        share=False,
        show_error=True
    )